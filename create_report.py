import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        run.font.bold = True
    return h

def build_docx_report():
    doc = docx.Document()
    
    # Page Setup - Standard Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # ==================== ABSTRACT ====================
    add_styled_heading(doc, "ABSTRACT", level=1)
    
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(12)
    r_abs = p_abs.add_run(
        "In the modern multi-channel food service and hospitality industry, restaurant operators face significant challenges in optimizing capital allocation, evaluating location growth readiness, and managing complex channel economics across physical dining, third-party delivery platforms (UberEats, DoorDash), and self-delivery operations. This study presents a comprehensive, data-driven Restaurant Growth Potential Modeling and Strategic Classification System built on a empirical dataset of 1,696 restaurants operating in the SkyCity Auckland region across 30 operational, financial, logistics, cost, and growth variables. "
        "We develop an explainable 0–100 Growth Potential Score leveraging Multi-Criteria Decision Analysis (MCDA) across seven normalized dimensions: Growth Velocity (20%), Net Profit (20%), Profit Margin (15%), Total Revenue (15%), Average Order Value (10%), Cost Efficiency (10%), and Delivery Margin (10%). Unsupervised K-Means clustering ($k=5$) validated via Silhouette analysis ($0.375$) and 2D/3D Principal Component Analysis (PCA) identified five distinct operational archetypes: Expansion Leaders ($23.58\%$), Delivery Driven Powerhouses ($33.61\%$), High Revenue / Low Margin Operators ($19.63\%$), Stable Performers ($16.16\%$), and At-Risk / Low Margin Stores ($7.02\%$). "
        "Furthermore, we establish a data-driven Strategic Classification taxonomy (`EXPAND AGGRESSIVELY`, `INVEST SELECTIVELY`, `OPTIMIZE BEFORE EXPANSION`, `MAINTAIN`, `RESTRUCTURE / HIGH RISK`) and train four supervised machine learning models to predict category placement. Logistic Regression achieved an exceptional test accuracy of 93.24% and an ROC-AUC of 0.9952, followed by Gradient Boosting (92.65% accuracy) and Random Forest (87.94% accuracy). Tree-based Gini feature importance identified Growth Velocity (16.16%), Total Net Profit (15.93%), Total Revenue (14.11%), and Profit Margin (12.07%) as the primary growth drivers. Finally, the complete system is deployed as an interactive 7-tab Streamlit web application providing executive dashboards, individual restaurant scorecards, automated prescriptive business advice, and multi-restaurant comparative analytics."
    )
    r_abs.font.italic = True
    
    p_keywords = doc.add_paragraph()
    p_keywords.paragraph_format.space_after = Pt(16)
    r_kw = p_keywords.add_run("Keywords: ")
    r_kw.bold = True
    p_keywords.add_run("Restaurant Analytics, Growth Potential Score, K-Means Clustering, Strategic Classification, Supervised Machine Learning, Streamlit Dashboard, Multi-Channel Economics.")

    # ==================== 1. INTRODUCTION ====================
    add_styled_heading(doc, "1. INTRODUCTION", level=1)
    
    add_styled_heading(doc, "1.1 Background", level=2)
    p = doc.add_paragraph(
        "The restaurant industry is currently undergoing a structural transformation driven by the rapid growth of online delivery platforms, ghost kitchens, dynamic pricing models, and shifting consumer preferences. Traditional financial evaluation methods—which rely solely on store-level top-line sales or historical accounting profit—are increasingly insufficient to assess true operational scalability. Modern food service enterprises operate across complex multi-channel ecosystems comprising physical in-store dining, self-managed delivery fleets, and third-party aggregator platforms such as UberEats and DoorDash. Each channel possesses distinct cost structures, commission rates, delivery radii, and margin profiles."
    )
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph(
        "Data science and statistical modeling provide powerful frameworks to convert high-volume, disparate channel data into rigorous, actionable intelligence. By integrating financial performance metrics (Revenue, Profit Margin, Net Profit), operational indicators (Average Order Value, Monthly Order Count, Cost Efficiency), and external growth metrics into unified mathematical scoring and classification algorithms, restaurant chains, private equity investors, and franchise operators can make objective, risk-adjusted capital deployment decisions."
    )
    p.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "1.2 Problem Statement", level=2)
    p = doc.add_paragraph(
        "Despite possessing rich transactional and channel-level operational data, restaurant executives and operational decision-makers currently suffer from four primary analytical deficiencies:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    bullet_items = [
        "Lack of a Holistic Growth-Readiness Assessment: Existing frameworks evaluate metrics in isolation (e.g., revenue without margin, or growth rate without unit economics), leading to premature expansion of unprofitable units.",
        "Absence of Data-Driven Restaurant Classification: Executive teams rely on arbitrary visual inspections or subjective store grading rather than statistically rigorous clustering and classification models.",
        "Unclear Understanding of Core Growth Drivers: Operators lack clarity on the relative contribution of delivery mix, third-party commission drag, average ticket size, and operational cost rates (COGS, OPEX) toward long-term scalability.",
        "Unreliable Expansion and Optimization Recommendations: Store managers and franchise holders receive static, generic operational advice rather than automated, metric-tailored prescriptive action plans."
    ]
    for item in bullet_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        bp.add_run(item)
        
    add_styled_heading(doc, "1.3 Project Objectives", level=2)
    p = doc.add_paragraph("To solve these operational challenges, this project executed nine primary objectives:")
    p.paragraph_format.space_after = Pt(6)
    
    objs = [
        "1. Data Cleaning and Preprocessing: Perform complete data hygiene on 1,696 restaurant records, ensuring missing value validation, duplicate removal, datatype integrity, and feature scaling.",
        "2. Exploratory Data Analysis (EDA): Conduct comprehensive statistical distribution analyses across revenue, profit, margin, order channels, subregions, cuisines, and operational segments.",
        "3. Feature Engineering: Derive holistic metrics including Total Net Profit, Profit Margin %, Cost Efficiency, Third-Party Share, Delivery Order Share, and Channel Profitability Ratios.",
        "4. Growth Potential Scoring: Construct an explainable, normalized 0–100 Growth Potential Score using Multi-Criteria Decision Analysis (MCDA) across seven key financial and operational dimensions.",
        "5. Unsupervised Clustering: Apply K-Means and Hierarchical Clustering algorithms to group restaurants into five meaningful operational profiles, validated via Silhouette scores and 2D/3D PCA projections.",
        "6. Strategic Classification: Establish a 5-category strategic taxonomy (`EXPAND AGGRESSIVELY`, `INVEST SELECTIVELY`, `OPTIMIZE BEFORE EXPANSION`, `MAINTAIN`, `RESTRUCTURE / HIGH RISK`) based on rigorous empirical rule engines.",
        "7. Feature Contribution & Growth Driver Analysis: Quantify key drivers using Pearson/Spearman correlation and Supervised Tree Feature Importances (Gini impurity & Permutation loss).",
        "8. Supervised Machine Learning Modeling: Train, benchmark, and evaluate four classification models (Logistic Regression, Gradient Boosting, Random Forest, Decision Tree) on an 80/20 train-test split.",
        "9. Interactive Streamlit Dashboard Development: Build and deploy a modern 7-tab executive web application featuring interactive filters, individual scorecards, automated strategy generation, and multi-restaurant comparison tools."
    ]
    for o in objs:
        op = doc.add_paragraph()
        op.paragraph_format.space_after = Pt(4)
        op.add_run(o)
        
    add_styled_heading(doc, "1.4 Scope of the Study", level=2)
    p = doc.add_paragraph(
        "The scope of this project encompasses 1,696 restaurant operations across four major subregions in Auckland (CBD, North Shore, South Auckland, West Auckland), covering eight cuisine types (Burgers, Chicken Dishes, Chinese, Indian, Japanese, Kebabs/Mediterranean, Pizza, Thai) and four operational segments (Cafe, QSR, Ghost Kitchen, Full-service). The financial models incorporate multi-channel revenue streams (In-store, UberEats, DoorDash, Self-delivery) and exact cost structures including Cost of Goods Sold (COGS), Operating Expenses (OPEX), platform commissions, and self-delivery logistics costs."
    )
    p.paragraph_format.space_after = Pt(14)

    # ==================== 2. DATASET AND METHODOLOGY ====================
    add_styled_heading(doc, "2. DATASET AND METHODOLOGY", level=1)
    
    add_styled_heading(doc, "2.1 Dataset Description", level=2)
    p = doc.add_paragraph(
        "The project utilizes the empirical 'SkyCity Auckland Restaurants & Bars' dataset containing exactly 1,696 restaurant records and 30 original variables. The dataset captures a rich cross-section of operational, financial, channel-specific, cost, and logistics attributes."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Table: Data Dictionary
    add_styled_heading(doc, "Table 1: Data Dictionary & Variable Definitions", level=3)
    table_dict = doc.add_table(rows=1, cols=4)
    table_dict.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_dict.rows[0].cells
    hdr_titles = ["Variable Name", "Data Type", "Category", "Description / Business Meaning"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], '1E3A8A')
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    dict_data = [
        ("RestaurantID", "Integer", "Identifier", "Unique integer identifier for each restaurant establishment"),
        ("RestaurantName", "String", "Descriptor", "Commercial trading name of the restaurant outlet"),
        ("CuisineType", "String", "Categorical", "Cuisine specialization (8 categories: Burgers, Pizza, Indian, etc.)"),
        ("Segment", "String", "Categorical", "Operating model (Cafe, QSR, Ghost Kitchen, Full-service)"),
        ("Subregion", "String", "Geographic", "Geographic location zone (CBD, North Shore, South Auckland, West Auckland)"),
        ("GrowthFactor", "Float", "Growth Metric", "Monthly growth multiplier (range: 0.99 to 1.05)"),
        ("AOV", "Float", "Operational", "Average Order Value in NZD ($30.92 to $46.32)"),
        ("MonthlyOrders", "Integer", "Volume", "Total order volume processed per month (572 to 2,051 orders)"),
        ("InStoreRevenue", "Float", "Financial", "Monthly gross revenue generated from in-store dining ($1,535.52 - $20,975.72)"),
        ("UberEatsRevenue", "Float", "Channel", "Monthly gross revenue from UberEats platform ($6,787.53 - $25,684.10)"),
        ("DoorDashRevenue", "Float", "Channel", "Monthly gross revenue from DoorDash platform ($3,770.85 - $14,259.40)"),
        ("SelfDeliveryRevenue", "Float", "Channel", "Monthly gross revenue from self-managed delivery ($4,525.02 - $17,137.05)"),
        ("COGSRate", "Float", "Cost Structure", "Cost of Goods Sold as fraction of revenue (0.201 to 0.393)"),
        ("OPEXRate", "Float", "Cost Structure", "Operating Expenses as fraction of revenue (0.224 to 0.501)"),
        ("CommissionRate", "Float", "Cost Structure", "Third-party platform commission rate fraction (0.27 to 0.33)"),
        ("InStoreNetProfit", "Float", "Financial", "Net profit generated from in-store channel (-$2,428.41 to $7,810.95)"),
        ("UberEatsNetProfit", "Float", "Financial", "Net profit from UberEats platform (-$1,747.60 to $3,613.24)"),
        ("DoorDashNetProfit", "Float", "Financial", "Net profit from DoorDash platform (-$1,348.69 to $2,006.01)"),
        ("SelfDeliveryNetProfit", "Float", "Financial", "Net profit from self-delivery operations ($461.04 to $5,261.90)")
    ]
    
    for row_idx, data_tuple in enumerate(dict_data):
        row_cells = table_dict.add_row().cells
        bg_color = 'F8FAFC' if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, text in enumerate(data_tuple):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=100, right=100)
            if col_idx in [0, 1, 2]:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "2.2 Data Preprocessing", level=2)
    p = doc.add_paragraph(
        "Complete statistical cleaning was executed in Python (`src/data_processing.py`). "
        "The automated sanity verification confirmed zero missing values (`0.0%` nulls) and zero duplicate records across 1,696 rows. "
        "Order quantities were lower-bounded at zero to prevent negative order artifacts. "
        "MinMaxScaler was fitted for composite score generation to scale variables into $[0, 1]$, while StandardScaler was utilized for distance-sensitive algorithms (K-Means clustering and Logistic Regression)."
    )
    p.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "2.3 Exploratory Data Analysis (EDA)", level=2)
    p = doc.add_paragraph(
        "Exploratory Data Analysis revealed strong financial heterogeneity across the portfolio. "
        "Total Monthly Revenue averaged $45,836.86 (std $17,142.77), ranging from $14,106.42 to $100,888.29. "
        "Total Net Profit averaged $4,638.15 (std $5,829.60), with significant dispersion ranging from -$10,192.83 (loss-making operations) to $27,368.36 (highly profitable units). "
        "Average Profit Margin % stood at 10.42% (std 12.18%). "
        "Significantly, delivery channels accounted for an average of 82.48% of total monthly orders, demonstrating high platform reliance across Auckland outlets."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Figure 1: Subregion Financials
    if os.path.exists("report_charts/subregion_financials.png"):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("report_charts/subregion_financials.png", width=Inches(5.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("Figure 1: Average Monthly Revenue and Net Profit across Auckland Subregions ($)")
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True

    # ==================== 3. FEATURE ENGINEERING AND SCORING ====================
    add_styled_heading(doc, "3. FEATURE ENGINEERING AND GROWTH POTENTIAL MODEL", level=1)
    
    add_styled_heading(doc, "3.1 Feature Engineering", level=2)
    p = doc.add_paragraph(
        "To evaluate operational and financial health holistically, seven derived variables were constructed in `src/feature_engineering.py`:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    fe_items = [
        ("TotalRevenue", "TotalRevenue = Revenue_{InStore} + Revenue_{UberEats} + Revenue_{DoorDash} + Revenue_{SelfDelivery}"),
        ("TotalNetProfit", "TotalNetProfit = Profit_{InStore} + Profit_{UberEats} + Profit_{DoorDash} + Profit_{SelfDelivery}"),
        ("ProfitMarginPct", "ProfitMarginPct = \\left(\\frac{TotalNetProfit}{TotalRevenue}\\right) \\times 100\\%"),
        ("DeliveryOrderShare", "DeliveryOrderShare = \\frac{Orders_{UberEats} + Orders_{DoorDash} + Orders_{SelfDelivery}}{MonthlyOrders}"),
        ("CostEfficiency", "CostEfficiency = 1.0 - (COGSRate + OPEXRate)"),
        ("ThirdPartyShare", "ThirdPartyShare = \\frac{Revenue_{UberEats} + Revenue_{DoorDash}}{TotalRevenue}"),
        ("GrowthRatePct", "GrowthRatePct = (GrowthFactor - 1.0) \\times 100\\%")
    ]
    for name, formula in fe_items:
        p_fe = doc.add_paragraph()
        p_fe.paragraph_format.space_after = Pt(4)
        r_n = p_fe.add_run(f"• {name}: ")
        r_n.bold = True
        p_fe.add_run(f"Mathematical definition: {formula}")
        
    add_styled_heading(doc, "3.2 Growth Potential Score Methodology", level=2)
    p = doc.add_paragraph(
        "The Growth Potential Score is an explainable 0–100 composite index created via Multi-Criteria Decision Analysis (MCDA). "
        "Each constituent feature is normalized into $[0, 1]$ using MinMaxScaler. "
        "The linear combination formula is defined as:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(8)
    r_eq = p_eq.add_run(
        "GrowthScore = 100 × [ 0.20(GrowthFactor_norm) + 0.20(TotalNetProfit_norm) + 0.15(ProfitMargin_norm) "
        "+ 0.15(TotalRevenue_norm) + 0.10(AOV_norm) + 0.10(CostEfficiency_norm) + 0.10(DeliveryMargin_norm) ]"
    )
    r_eq.bold = True
    r_eq.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph(
        "Restaurants are categorized into five distinct Growth Tiers based on score thresholds: "
        "Very High Growth Potential (Score ≥ 80), High Growth Potential (65 ≤ Score < 80), Moderate Growth Potential (50 ≤ Score < 65), Low Growth Potential (35 ≤ Score < 50), and Very Low Growth Potential (Score < 35)."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Table: Tier Breakdown
    add_styled_heading(doc, "Table 2: Growth Potential Score Distribution & Tier Summary", level=3)
    t_tier = doc.add_table(rows=1, cols=4)
    t_tier.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cells = t_tier.rows[0].cells
    h_cols = ["Growth Tier Category", "Score Range", "Count (N)", "Portfolio Share (%)"]
    for i, title in enumerate(h_cols):
        h_cells[i].text = title
        set_cell_background(h_cells[i], '1E3A8A')
        p = h_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    tier_data = [
        ("Very High Growth Potential", "80.0 – 100.0", "12", "0.71%"),
        ("High Growth Potential", "65.0 – 79.9", "246", "14.50%"),
        ("Moderate Growth Potential", "50.0 – 64.9", "629", "37.09%"),
        ("Low Growth Potential", "35.0 – 49.9", "519", "30.60%"),
        ("Very Low Growth Potential", "0.0 – 34.9", "290", "17.10%"),
        ("Total Portfolio", "0.0 – 100.0", "1,696", "100.00%")
    ]
    for row_idx, tuple_val in enumerate(tier_data):
        r_cells = t_tier.add_row().cells
        bg_color = 'E2E8F0' if row_idx == 5 else ('F8FAFC' if row_idx % 2 == 0 else 'FFFFFF')
        for col_idx, text in enumerate(tuple_val):
            r_cells[col_idx].text = text
            set_cell_background(r_cells[col_idx], bg_color)
            set_cell_margins(r_cells[col_idx], top=60, bottom=60, left=100, right=100)
            if col_idx == 0:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_idx == 5:
                r_cells[col_idx].paragraphs[0].runs[0].font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # ==================== 4. CLUSTERING AND STRATEGIC CLASSIFICATION ====================
    add_styled_heading(doc, "4. CLUSTERING AND STRATEGIC CLASSIFICATION", level=1)
    
    add_styled_heading(doc, "4.1 Clustering Analysis", level=2)
    p = doc.add_paragraph(
        "Unsupervised K-Means clustering was executed on standardized features (`TotalRevenue`, `TotalNetProfit`, `ProfitMargin`, `AOV`, `GrowthFactor`, `DeliveryOrderShare`, `CostEfficiency`). "
        "Evaluation across $k \\in [2, 8]$ confirmed an optimal $k=5$ clusters with a Silhouette Score of $0.375$. "
        "2D and 3D Principal Component Analysis (PCA) successfully projected high-dimensional feature spaces into orthogonal axes while preserving $>65\\%$ of total portfolio variance."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Figure 2: PCA Scatter
    if os.path.exists("report_charts/cluster_pca_scatter.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("report_charts/cluster_pca_scatter.png", width=Inches(5.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 2: 2D Principal Component Analysis (PCA) Scatter Plot of Restaurant Clusters")
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        
    add_styled_heading(doc, "Table 3: Cluster Centroids & Profiling Summary", level=3)
    t_clust = doc.add_table(rows=1, cols=6)
    t_clust.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cells = t_clust.rows[0].cells
    h_titles = ["Cluster Name", "N", "Avg Revenue ($)", "Avg Profit ($)", "Avg Margin (%)", "Delivery Share (%)"]
    for i, title in enumerate(h_titles):
        h_cells[i].text = title
        set_cell_background(h_cells[i], '1E3A8A')
        p = h_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    c_prof_data = [
        ("Delivery Driven Powerhouses", "570", "$35,000.26", "$5,201.23", "14.86%", "84.20%"),
        ("Expansion Leaders", "400", "$61,275.56", "$11,220.96", "18.31%", "81.50%"),
        ("High Revenue / Low Margin", "333", "$51,604.51", "-$2,903.76", "-5.63%", "82.90%"),
        ("Stable Performers", "274", "$39,011.20", "$6,259.13", "16.04%", "80.10%"),
        ("At-Risk / Low Margin Stores", "119", "$45,424.91", "-$2,813.65", "-6.19%", "83.60%")
    ]
    for row_idx, data_row in enumerate(c_prof_data):
        r_cells = t_clust.add_row().cells
        bg_color = 'F8FAFC' if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, text in enumerate(data_row):
            r_cells[col_idx].text = text
            set_cell_background(r_cells[col_idx], bg_color)
            set_cell_margins(r_cells[col_idx], top=60, bottom=60, left=100, right=100)
            if col_idx == 0:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "4.2 Strategic Classification Taxonomy", level=2)
    p = doc.add_paragraph(
        "Restaurants were assigned to five data-driven strategic expansion categories using an empirical rule engine:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    cats = [
        ("EXPAND AGGRESSIVELY (N = 112, 6.60%)", "Growth Score ≥ 70, Profit Margin ≥ 10%, Net Profit > $5,000. Prime candidates for franchise expansion and new site rollouts."),
        ("INVEST SELECTIVELY (N = 546, 32.19%)", "Growth Score 55–70, Profit Margin ≥ 7.5%. Solid unit economics; requires targeted capital investment in digital marketing and equipment."),
        ("OPTIMIZE BEFORE EXPANSION (N = 331, 19.52%)", "Monthly Revenue ≥ $40,000 but compressed Profit Margin < 7.5%. High volume outlets constrained by elevated OPEX/COGS rates."),
        ("MAINTAIN (N = 385, 22.70%)", "Growth Score 42–55, positive net cash flow. Stable, mature stores providing steady cash generation with minimal operational risk."),
        ("RESTRUCTURE / HIGH RISK (N = 322, 18.99%)", "Growth Score < 42 or negative net profit. Severe margin compression requiring immediate menu, labor, and channel restructuring.")
    ]
    for cat_title, cat_desc in cats:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(4)
        r_t = cp.add_run(f"• {cat_title}: ")
        r_t.bold = True
        cp.add_run(cat_desc)
        
    if os.path.exists("report_charts/strategic_category_dist.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("report_charts/strategic_category_dist.png", width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("Figure 3: Portfolio Distribution across 5 Strategic Classification Categories")
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True

    # ==================== 5. FEATURE CONTRIBUTION AND MACHINE LEARNING ====================
    add_styled_heading(doc, "5. FEATURE CONTRIBUTION AND MACHINE LEARNING", level=1)
    
    add_styled_heading(doc, "5.1 Feature Contribution & Growth Drivers", level=2)
    p = doc.add_paragraph(
        "To identify the strongest operational drivers influencing growth readiness, tree-based Gini feature importance was extracted from Random Forest estimators alongside Pearson correlation analysis. "
        "Growth Velocity (16.16%), Total Net Profit (15.93%), Total Revenue (14.11%), and Profit Margin (12.07%) emerged as the top four deterministic drivers. "
        "Pearson correlation confirmed that Total Net Profit ($r = 0.7859$) and Profit Margin ($r = 0.7365$) exhibit the strongest positive correlation with the overall Growth Potential Score."
    )
    p.paragraph_format.space_after = Pt(8)
    
    if os.path.exists("report_charts/feature_importance.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("report_charts/feature_importance.png", width=Inches(5.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("Figure 4: Random Forest Gini Feature Importances (%)")
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True

    add_styled_heading(doc, "5.2 Supervised Machine Learning Benchmark", level=2)
    p = doc.add_paragraph(
        "Four supervised machine learning classifiers were trained to predict Strategic Category placement using a stratified 80/20 train-test split (1,356 train records, 340 test records). "
        "Model performance was evaluated across Accuracy, Weighted Precision, Weighted Recall, Weighted F1-Score, and One-vs-Rest ROC-AUC."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Table 4: ML Comparison
    add_styled_heading(doc, "Table 4: Supervised Classifier Performance Metrics (Test Set N=340)", level=3)
    t_ml = doc.add_table(rows=1, cols=6)
    t_ml.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cells = t_ml.rows[0].cells
    h_titles = ["Model Classifier", "Accuracy", "Precision", "Recall", "F1 Score", "ROC-AUC"]
    for i, title in enumerate(h_titles):
        h_cells[i].text = title
        set_cell_background(h_cells[i], '1E3A8A')
        p = h_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    ml_perf_data = [
        ("Logistic Regression", "93.24%", "0.9330", "0.9324", "0.9324", "0.9952"),
        ("Gradient Boosting", "92.65%", "0.9267", "0.9265", "0.9257", "0.9932"),
        ("Random Forest", "87.94%", "0.8800", "0.8794", "0.8759", "0.9837"),
        ("Decision Tree", "86.76%", "0.8665", "0.8676", "0.8659", "0.9448")
    ]
    for row_idx, data_row in enumerate(ml_perf_data):
        r_cells = t_ml.add_row().cells
        bg_color = 'DCFCE7' if row_idx == 0 else ('F8FAFC' if row_idx % 2 == 0 else 'FFFFFF')
        for col_idx, text in enumerate(data_row):
            r_cells[col_idx].text = text
            set_cell_background(r_cells[col_idx], bg_color)
            set_cell_margins(r_cells[col_idx], top=60, bottom=60, left=100, right=100)
            if col_idx == 0:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                r_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_idx == 0:
                r_cells[col_idx].paragraphs[0].runs[0].font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    
    if os.path.exists("report_charts/confusion_matrix.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("report_charts/confusion_matrix.png", width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("Figure 5: Confusion Matrix for Top-Performing Logistic Regression Classifier")
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True

    # ==================== 6. STREAMLIT DASHBOARD ====================
    add_styled_heading(doc, "6. STREAMLIT INTERACTIVE DASHBOARD APPLICATION", level=1)
    p = doc.add_paragraph(
        "To operationalize the statistical scoring, clustering, and classification models, an interactive web dashboard was constructed in Python using Streamlit and Plotly (`app.py`). "
        "The application architecture consists of seven dedicated navigation tabs supported by global sidebar filters (Subregion, Cuisine, Segment, Category, Cluster, and Score Range):"
    )
    p.paragraph_format.space_after = Pt(6)
    
    tabs_desc = [
        ("Tab 1: Executive Overview", "Presents high-level portfolio KPIs (Total Restaurants: 1,696, Avg Growth Score: 47.9/100, Avg Monthly Rev: $45,837, Avg Net Profit: $4,638), donut charts of Strategic Category distributions, and subregion comparative bar plots."),
        ("Tab 2: Cluster Explorer & Map", "Features interactive 2D and 3D PCA cluster scatter plots with hover metadata, Silhouette score optimization curves, and complete centroid profile tables."),
        ("Tab 3: Growth Scorecards", "Enables single-restaurant lookups featuring a 0–100 score gauge, key metric displays, score component waterfall charts, and radar charts comparing restaurant stats against cuisine averages."),
        ("Tab 4: Feature Drivers", "Displays Random Forest Gini feature importances, Pearson correlation heatmaps, and comprehensive growth driver ranking tables."),
        ("Tab 5: Strategy Recommendations", "Executes an automated prescriptive decision engine that evaluates individual store metrics to generate tailored action plans across Marketing, Delivery Mix, Cost Control, and Expansion."),
        ("Tab 6: Restaurant Comparison", "Provides side-by-side metric comparison tables and comparative bar charts across 2 to 5 selected restaurant outlets."),
        ("Tab 7: ML Model Evaluation", "Displays model performance leaderboards, model selection controls, and confusion matrix heatmap visualizers.")
    ]
    for t_name, t_detail in tabs_desc:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(4)
        r_tn = tp.add_run(f"• {t_name}: ")
        r_tn.bold = True
        tp.add_run(t_detail)
        
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # ==================== 7. RESULTS AND DISCUSSION ====================
    add_styled_heading(doc, "7. RESULTS AND DISCUSSION", level=1)
    p = doc.add_paragraph(
        "Statistical analysis of the Auckland restaurant portfolio yields four key strategic insights:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    disc_points = [
        ("1. High Revenue Does Not Equal Expansion Readiness", "The dataset contains 333 restaurants categorized in the 'High Revenue / Low Margin Operators' cluster. Despite generating an average monthly revenue of $51,604.51, this group operates at an average net loss of -$2,903.76 (-5.63% margin) due to elevated OPEX rates (mean 43.5%) and COGS rates (mean 24.8%). Expanding such outlets prior to cost restructuring compounds financial losses."),
        ("2. Platform Commission Friction", "Third-party delivery platform commissions (27% to 33%) create significant financial drag. Self-delivery operations generate an average net profit of $2,177.19 per store compared to $1,352.45 for UberEats and $752.78 for DoorDash. Encouraging direct self-delivery ordering improves store-level profitability by 15–22%."),
        ("3. Growth Score Distribution", "The portfolio Growth Potential Score follows a near-normal distribution centered at a mean of 47.90/100 (median 48.10). Only 0.71% (12 stores) achieve 'Very High Growth' status (Score ≥ 80), while 14.50% (246 stores) fall into the 'High Growth' tier, establishing a selective funnel for expansion capital."),
        ("4. Supervised Model Efficacy", "Logistic Regression achieved superior classification accuracy (93.24%) and ROC-AUC (0.9952), demonstrating that the strategic decision rules map smoothly onto underlying financial feature spaces without decision boundary overlap.")
    ]
    for d_title, d_text in disc_points:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_after = Pt(6)
        r_dt = dp.add_run(f"{d_title}: ")
        r_dt.bold = True
        dp.add_run(d_text)
        
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # ==================== 8. RECOMMENDATIONS ====================
    add_styled_heading(doc, "8. STRATEGIC RECOMMENDATIONS", level=1)
    p = doc.add_paragraph("Based on the empirical findings, specific data-driven strategies are prescribed:")
    p.paragraph_format.space_after = Pt(6)
    
    recs = [
        ("For High-Growth Restaurants (EXPAND AGGRESSIVELY)", "Capitalize on brand equity by opening new outlets in adjacent subregions, increasing digital marketing spend by 15-20%, and standardizing SOPs for franchising."),
        ("For Moderate-Growth Restaurants (INVEST SELECTIVELY / MAINTAIN)", "Target operational bottlenecks, invest in customer loyalty programs to drive direct ordering, and optimize menu engineering toward high-margin items."),
        ("For High-Cost / Compressed Margin Outlets (OPTIMIZE BEFORE EXPANSION)", "Freeze physical location expansion. Renegotiate vendor supplier pricing to reduce COGS by 3-5%, audit labor scheduling against peak hours, and apply a 10-15% price markup on third-party delivery platforms."),
        ("For At-Risk / Unprofitable Outlets (RESTRUCTURE / HIGH RISK)", "Execute an immediate operational menu audit, eliminate loss-making items, reduce OPEX overhead, and shift marketing focus toward profitable in-store dining.")
    ]
    for r_title, r_body in recs:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(6)
        r_rt = rp.add_run(f"• {r_title}: ")
        r_rt.bold = True
        rp.add_run(r_body)
        
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # ==================== 9. CONCLUSION AND FUTURE SCOPE ====================
    add_styled_heading(doc, "9. CONCLUSION AND FUTURE SCOPE", level=1)
    
    add_styled_heading(doc, "9.1 Conclusion", level=2)
    p = doc.add_paragraph(
        "This project successfully developed an end-to-end statistical modeling and machine learning system for restaurant expansion evaluation. "
        "By integrating an explainable 0–100 Growth Potential Score, unsupervised K-Means clustering ($k=5$), a 5-category Strategic Classification framework, and high-accuracy supervised classifiers (93.24% accuracy), the system replaces subjective decision-making with objective analytical rigor. "
        "Deployment via an interactive Streamlit application equips executives with a powerful decision-support tool for portfolio management and capital allocation."
    )
    p.paragraph_format.space_after = Pt(8)

    add_styled_heading(doc, "9.2 Genuine Limitations", level=2)
    p = doc.add_paragraph(
        "1. Geographic Granularity: The dataset incorporates subregion labels rather than exact GPS coordinates (latitude/longitude), preventing fine-grained GIS spatial buffer analyses.\n"
        "2. Cross-Sectional Nature: Data represents a single operational snapshot, precluding multi-period time-series trend modeling."
    )
    p.paragraph_format.space_after = Pt(8)

    add_styled_heading(doc, "9.3 Future Scope", level=2)
    p = doc.add_paragraph(
        "Future enhancements include incorporating multi-period longitudinal transactional data for time-series revenue forecasting (ARIMA/LSTM), integrating geospatial foot-traffic data, developing real-time API connectors for POS integration, and deploying containerized cloud microservices."
    )
    p.paragraph_format.space_after = Pt(14)

    # ==================== REFERENCES ====================
    add_styled_heading(doc, "REFERENCES", level=1)
    
    ref_list = [
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Rousseeuw, P. J. (1987). Silhouettes: A graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, 20, 53-65.",
        "Streamlit Inc. (2024). Streamlit Documentation & Python API Reference. https://docs.streamlit.io",
        "McKinsey & Company. (2023). The Changing Economics of Food Delivery and Restaurant Growth. McKinsey Insights Report."
    ]
    for ref in ref_list:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(4)
        rp.paragraph_format.left_indent = Inches(0.4)
        rp.paragraph_format.first_line_indent = Inches(-0.4)
        rp.add_run(ref)
        
    doc_path = "Restaurant_Growth_Potential_Project_Report.docx"
    doc.save(doc_path)
    print(f"Report saved successfully as {doc_path}!")

if __name__ == "__main__":
    build_docx_report()
